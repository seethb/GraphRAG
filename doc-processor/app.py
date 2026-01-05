from flask import Flask, request, jsonify
from flask_cors import CORS
import PyPDF2
import docx
import io

app = Flask(__name__)
CORS(app)

@app.route('/health', methods=['GET'])
def health():
    return jsonify({'status': 'ok'})

@app.route('/extract/pdf', methods=['POST'])
def extract_pdf():
    """Extract text from PDF"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
            
        file = request.files['file']
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file.read()))
        
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n\n"
        
        return jsonify({
            'text': text,
            'pages': len(pdf_reader.pages),
            'char_count': len(text)
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/extract/docx', methods=['POST'])
def extract_docx():
    """Extract text from Word document"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
            
        file = request.files['file']
        doc = docx.Document(io.BytesIO(file.read()))
        
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n\n"
        
        return jsonify({
            'text': text,
            'paragraphs': len(doc.paragraphs),
            'char_count': len(text)
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/extract/text', methods=['POST'])
def extract_text():
    """Extract text from plain text file"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
            
        file = request.files['file']
        text = file.read().decode('utf-8')
        
        return jsonify({
            'text': text,
            'char_count': len(text)
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5006, debug=True)
